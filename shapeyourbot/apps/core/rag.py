from langchain_community.document_loaders import PDFPlumberLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_ollama import OllamaEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_ollama.llms import OllamaLLM
from docx.api import Document
from mrkdwn_analysis import MarkdownAnalyzer

PROMPT_TEMPLATE = """
You are a helpful assistant with the goal to provide an answer based on the context. Only use the provided context to answer the query. 
If unsure, state that you don't know. Answer in the language that the query is asked in.

Query: {user_query} 
Context: {document_context} 
Answer:
"""

EMBEDDING_MODEL = OllamaEmbeddings(model="deepseek-r1:1.5b")
DOCUMENT_VECTOR_DB = InMemoryVectorStore(EMBEDDING_MODEL)
LANGUAGE_MODEL = OllamaLLM(model="deepseek-r1:1.5b")

def receive_llm_answer(user_query):
    relevant_docs = find_related_documents(user_query)
    ai_response = generate_answer(user_query, relevant_docs)
    print(ai_response)
    return ai_response

def index_documents(document_chunks):
    DOCUMENT_VECTOR_DB.add_documents(document_chunks)

def find_related_documents(query):
    return DOCUMENT_VECTOR_DB.similarity_search(query)

def generate_answer(user_query, context_documents):
    context_text = "\n\n".join([doc["page_content"] for doc in context_documents])
    conversation_prompt = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
    response_chain = conversation_prompt | LANGUAGE_MODEL
    return response_chain.invoke({"user_query": user_query, "document_context": context_text})

def chunk_documents(raw_docs, format):
    text_processor = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        add_start_index=True
    )

    if format == "pdf":
        return text_processor.split_documents(raw_docs)
    elif format == "docx":
        raw_docs = text_processor.create_documents(raw_docs)
        chunks = text_processor.split_documents(raw_docs)
        return chunks
    elif format == "md":
        raw_docs = text_processor.create_documents(raw_docs)
        chunks = text_processor.split_documents(raw_docs)
        return chunks
    elif format == "txt":
        raw_docs = text_processor.create_documents(raw_docs)
        chunks = text_processor.split_documents(raw_docs)
        return chunks

def process_document(document_path):
    if document_path.endswith('.pdf'):
        raw_docs = extract_pdf_documents(document_path)
        processed_chunks = chunk_documents(raw_docs, "pdf")
    elif document_path.endswith('.docx'):
        raw_docs = extract_docx_document(document_path)
        processed_chunks = chunk_documents(raw_docs, "docx")
    elif document_path.endswith('.md'):
        raw_docs = extract_md_document(document_path)
        processed_chunks = chunk_documents(raw_docs, "md")
    elif document_path.endswith('.txt'):
        raw_docs = extract_txt_document(document_path)
        processed_chunks = chunk_documents(raw_docs, "txt")
    else:
        print("Unsupported document format.")

    index_documents(processed_chunks)
    print("Document processed and indexed.")

def extract_txt_document(document_path):
    with open(document_path, "r") as file:
        combined_content = file.read()
        print("combined_content", combined_content)
    return combined_content

# currently only extracting paragraphs, headers and links from md files
def extract_md_document(document_path):
    analyzer = MarkdownAnalyzer(document_path)
    combined_content = ""

    paragraphs = analyzer.identify_paragraphs()
    if paragraphs:
        paragraphs = '\n'.join(paragraphs['Paragraph'])
        combined_content += (paragraphs + "\n")

    headers = analyzer.identify_headers()
    if headers:
        headers = '\n'.join(headers['Headers'])
        combined_content += (headers + "\n")

    links = analyzer.identify_links()
    if links:
        links = '\n'.join(links['Links'])
        combined_content += (links + "\n")
    return combined_content

def extract_docx_document(document_path):
    document = Document(document_path)
    text = extract_docx_text(document)
    tables = extract_docx_tables(document)
    combined_content = text + "\n" + tables
    return combined_content

def extract_pdf_documents(file_path):
    document_loader = PDFPlumberLoader(file_path)
    return document_loader.load()

def extract_docx_tables(document):
    tables = ""
    count = 0

    for table in document.tables:
        count += 1
        table_content = f"#{count} Table\n"

        for row in table.rows:
            table_content += "|".join([cell.text for cell in row.cells])
            table_content += "\n"

        table_content += "\n"
        tables += table_content

    return tables

def extract_docx_text(document):
    all_text = ""

    for p in document.paragraphs:
        all_text += p.text
        all_text += "\n"

    return all_text